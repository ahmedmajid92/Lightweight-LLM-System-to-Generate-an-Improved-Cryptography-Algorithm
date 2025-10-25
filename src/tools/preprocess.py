import docx
import json
import os

def preprocess(input_docx: str, output_json: str):
    doc = docx.Document(input_docx)
    # 1) Parse characteristics table
    char_table = doc.tables[0]
    headers_char = [cell.text.strip() for cell in char_table.rows[0].cells]
    data_char = []
    for row in char_table.rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        data_char.append(dict(zip(headers_char, values)))

    # 2) Parse usage table
    usage_table = doc.tables[1]
    headers_usage = [cell.text.strip() for cell in usage_table.rows[0].cells]
    data_usage = []
    for row in usage_table.rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        data_usage.append(dict(zip(headers_usage, values)))

    # 3) Merge and extract pseudocode
    algos = []
    for entry in data_char:
        name = entry['Algorithm']
        usage = next((u for u in data_usage if u['Algorithm'] == name), {})
        # Extract pseudocode paragraphs
        capture = False
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text.startswith(f"Algorithm {name}"):
                capture = True
            elif text.startswith("Algorithm ") and capture:
                break
            if capture and text:
                lines.append(text)
        entry['pseudocode'] = "\n".join(lines)
        entry.update(usage)
        algos.append(entry)

    os.makedirs(os.path.dirname(output_json), exist_ok=True)
    with open(output_json, 'w') as f:
        json.dump(algos, f, indent=4)

if __name__ == '__main__':
    preprocess('data/BlockCipherTable.docx', 'data/algorithms.json')